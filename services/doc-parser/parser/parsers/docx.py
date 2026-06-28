"""DOCX document parser using python-docx."""

import importlib.util
import io
import time
from typing import Any

import structlog

from parser.models import DocumentFormat, ImageRef, ParsedDocument, TableData
from parser.parsers.base import BaseParser

logger = structlog.get_logger()


class DocxParser(BaseParser):
    """Parser for DOCX documents using python-docx."""

    @property
    def format(self) -> DocumentFormat:
        return DocumentFormat.DOCX

    def is_available(self) -> bool:
        """Check if python-docx is available."""
        return importlib.util.find_spec("docx") is not None

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse DOCX content and extract text, tables, and metadata."""
        from docx import Document

        start_time = time.time()
        text_parts: list[str] = []
        tables: list[TableData] = []
        images: list[ImageRef] = []
        metadata: dict[str, Any] = {}

        try:
            doc = Document(io.BytesIO(content))

            # Extract core properties (metadata)
            core_props = doc.core_properties
            if core_props:
                metadata = {
                    "title": core_props.title,
                    "author": core_props.author,
                    "subject": core_props.subject,
                    "keywords": core_props.keywords,
                    "created": str(core_props.created) if core_props.created else None,
                    "modified": str(core_props.modified) if core_props.modified else None,
                }
                # Remove None values
                metadata = {k: v for k, v in metadata.items() if v}

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                rows_data: list[list[str]] = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows_data.append(row_data)

                if rows_data:
                    headers = rows_data[0] if rows_data else []
                    data_rows = rows_data[1:] if len(rows_data) > 1 else []
                    tables.append(TableData(headers=headers, rows=data_rows))
                    table_text = table_rows_to_retrieval_text(headers, data_rows)
                    if table_text:
                        text_parts.append(table_text)

            # Count images
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    images.append(ImageRef(index=image_count))
                    image_count += 1

        except Exception as e:
            logger.error("docx_parse_error", error=str(e), filename=filename)
            raise

        parse_time_ms = (time.time() - start_time) * 1000

        logger.info(
            "docx_parsed",
            filename=filename,
            paragraphs=len(text_parts),
            tables=len(tables),
            images=len(images),
            parse_time_ms=round(parse_time_ms, 2),
        )

        return ParsedDocument(
            text="\n\n".join(text_parts),
            format=DocumentFormat.DOCX,
            page_count=0,  # DOCX doesn't have explicit pages
            metadata=metadata,
            tables=tables,
            images=images,
            parse_time_ms=parse_time_ms,
        )


def table_rows_to_retrieval_text(headers: list[str], rows: list[list[str]]) -> str:
    """Render table content as header/value text for retrieval."""
    parts: list[str] = []
    header_text = " ".join(cell.strip() for cell in headers if cell.strip())
    if header_text:
        parts.append(header_text)

    for row in rows:
        row_parts: list[str] = []
        for idx, value in enumerate(row):
            value = value.strip()
            if not value:
                continue
            header = headers[idx].strip() if idx < len(headers) else ""
            row_parts.append(f"{header} {value}" if header else value)
        if row_parts:
            parts.append(" ".join(row_parts))

    return "\n".join(parts)
