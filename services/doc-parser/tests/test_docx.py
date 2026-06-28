"""Tests for the DOCX parser."""

import io

from docx import Document

from parser.parsers.docx import DocxParser


def docx_with_contract_table() -> bytes:
    """Build a DOCX whose retrievable content lives only in a table."""
    document = Document()
    table = document.add_table(rows=2, cols=3)
    headers = ["customer", "year", "contract_amount"]
    values = ["HGC", "2025", "1200"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for idx, value in enumerate(values):
        table.cell(1, idx).text = value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_table_content_is_included_in_retrieval_text():
    parser = DocxParser()

    result = parser.parse(docx_with_contract_table(), "contract.docx")

    assert "customer HGC" in result.text
    assert "year 2025" in result.text
    assert "contract_amount 1200" in result.text
    assert result.tables[0].headers == ["customer", "year", "contract_amount"]
